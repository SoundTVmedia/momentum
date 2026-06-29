#!/usr/bin/env python3
"""Generate native-ios-effort-estimate.docx — React Native iOS migration effort assessment."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "docs/native-ios-effort-estimate.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def main():
    doc = Document()

    title = doc.add_heading("Native iOS / React Native Migration — Effort Estimate", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Document purpose: ").bold = True
    p.add_run(
        "Assess the level of effort to migrate the current Feedback (Momentum) codebase "
        "and data structure into a fully native iOS experience built with React Native."
    )
    p = doc.add_paragraph()
    p.add_run("Last updated: ").bold = True
    p.add_run("June 26, 2026")

    doc.add_paragraph()

    # Executive summary
    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        "This is a medium-to-large rewrite, not a conversion. The backend and data model are "
        "largely reusable, but the iOS client UI and device integrations would need to be rebuilt.",
    )
    add_para(
        doc,
        "Estimated effort ranges (one experienced iOS/React Native engineer, part-time QA):",
    )
    add_table(
        doc,
        ["Scope", "Timeline"],
        [
            ("Prototype / thin native shell", "2–4 weeks"),
            ("Usable MVP (feed, auth, profiles, clip viewing, basic upload)", "6–10 weeks"),
            ("Feature-parity iOS app", "12–20+ weeks"),
            ("Swift/SwiftUI rewrite (not React Native)", "16–30+ weeks"),
        ],
    )

    # Current architecture
    add_heading(doc, "Current Architecture", 1)
    add_bullets(
        doc,
        [
            "React 19 web app (Vite) with react-router — ~37 pages, 120+ components",
            "Capacitor 7 iOS shell wrapping dist/ (or remote Worker URL) — not a native UI layer",
            "Cloudflare Worker API (Hono) with D1 database, R2 storage, Cloudflare Stream",
            "Session-cookie auth via @getmocha/users-service",
            "Native bridges today: push notifications, filesystem, gallery save (Capacitor plugins)",
            "Heavy browser APIs: MediaRecorder, getUserMedia, IndexedDB upload outbox, canvas thumbnails",
        ],
    )

    # What is reusable
    add_heading(doc, "What Is Reusable", 1)
    add_para(doc, "These layers can largely stay as-is or be shared with minimal changes:")
    add_bullets(
        doc,
        [
            "Cloudflare Worker API and all /api/* endpoints",
            "D1 schema and 60+ SQL migrations (clips, users, shows, notifications, gamification, etc.)",
            "R2 / Stream video pipeline and multipart upload sessions",
            "Third-party integrations: JamBase, Ticketmaster, Stripe, AudD/ACRCloud, YouTube",
            "Shared TypeScript business logic in src/shared/ (show matching, content feed, types)",
            "API contracts and data shapes consumed by the client",
        ],
    )

    # What must be rewritten
    add_heading(doc, "What Must Be Rewritten", 1)
    add_bullets(
        doc,
        [
            "All UI: DOM components → React Native (View, Text, FlatList, etc.)",
            "Routing: react-router → React Navigation (stack, tabs, deep links)",
            "Styling: Tailwind/CSS → StyleSheet or NativeWind (no direct transfer)",
            "Video capture & upload: UploadClip.tsx (~3,300 lines), QuickCaptureOverlay, camera zoom, thumbnails",
            "Upload outbox: IndexedDB + in-memory blob registry → AsyncStorage/SQLite + filesystem",
            "Auth/OAuth: cookie sessions and /auth/callback → native deep links + secure token storage",
            "Push notifications: Capacitor plugin → APNs + React Native push library",
            "Photo library, camera, mic permissions → native modules (expo-camera, react-native-vision-camera, etc.)",
            "Background upload: scaffolded in native-bridge.ts but not fully wired — needs URLSession implementation",
            "Live features: WebSocket chat, live broadcast, polls — native equivalents",
            "Admin/moderation dashboards, analytics charts (recharts) — lower priority, significant UI work",
        ],
    )

    # High-risk areas
    add_heading(doc, "High-Risk / High-Effort Areas", 1)
    add_table(
        doc,
        ["Area", "Why It's Hard", "Rough Effort"],
        [
            (
                "Clip capture & upload",
                "MediaRecorder, blob lifecycle, offline queue, multipart resume, song ID, show matching",
                "4–8 weeks",
            ),
            (
                "Auth & sessions",
                "Cookie-based Mocha auth, OAuth redirect flow, password reset deep links",
                "1–2 weeks",
            ),
            (
                "Video playback feed",
                "HLS via hls.js today; need expo-av or react-native-video + prefetch/limiter logic",
                "2–3 weeks",
            ),
            (
                "Push & notifications",
                "APNs registration, badge counts, in-app notification panel",
                "1–2 weeks",
            ),
            (
                "Discover & search",
                "JamBase geo search, trending carousels, infinite scroll grids",
                "2–4 weeks",
            ),
            (
                "Admin / superadmin",
                "Moderation panels, role management, analytics — web-only tooling today",
                "3–5 weeks (optional)",
            ),
        ],
    )

    # Recommended approach
    add_heading(doc, "Recommended Phased Approach", 1)
    add_para(doc, "Keep the existing Worker API and shared logic. Build a React Native iOS app alongside the web app.")
    add_bullets(
        doc,
        [
            "Phase 1 — Foundation (2–3 weeks): RN project scaffold, API client, auth, navigation shell",
            "Phase 2 — Core experience (3–4 weeks): Home feed, clip playback, profiles, saved/liked clips",
            "Phase 3 — Capture & upload (4–6 weeks): Camera, upload queue, offline outbox, gallery save",
            "Phase 4 — Discovery (2–3 weeks): Discover, nearby/tonight shows, artist/venue pages",
            "Phase 5 — Engagement (2–3 weeks): Push, notifications, show marks, gamification/points",
            "Phase 6 — Parity (ongoing): Live features, admin, sponsor/ambassador dashboards, analytics",
        ],
    )

    # Alternatives
    add_heading(doc, "Alternatives Considered", 1)
    add_bullets(
        doc,
        [
            "Stay on Capacitor: lowest effort; current path. Limitations: WebView performance, background upload gaps.",
            "Capacitor + native plugins for capture/upload only: hybrid — keep web UI, replace hardest native pieces.",
            "React Native (recommended for 'native React iOS'): best balance of code reuse (shared TS logic) and native UX.",
            "Swift/SwiftUI: maximum native feel; zero UI reuse; longest timeline unless team is iOS-first.",
        ],
    )

    # Key files reference
    add_heading(doc, "Key Codebase References", 1)
    add_table(
        doc,
        ["Path", "Role"],
        [
            ("src/react-app/App.tsx", "37 routes, providers, native bootstrap"),
            ("src/react-app/pages/UploadClip.tsx", "Largest client file — capture, tagging, upload enqueue"),
            ("src/react-app/contexts/ClipUploadQueueContext.tsx", "FIFO upload queue, IDB recovery, retry logic"),
            ("src/react-app/lib/native-bridge.ts", "Capacitor push, filesystem, gallery save"),
            ("src/react-app/lib/apiFetch.ts", "Session-cookie API client"),
            ("src/worker/index.ts", "Main Worker API surface"),
            ("migrations/", "D1 schema (60+ migrations)"),
            ("capacitor.config.ts", "Current iOS shell config"),
        ],
    )

    add_para(
        doc,
        "See also: docs/upload-flow-developer-summary.docx (upload pipeline), "
        "docs/upload-outbox-build.md (outbox architecture).",
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("Feedback (Momentum) — Native iOS Effort Estimate")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
