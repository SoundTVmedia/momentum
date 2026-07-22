#!/usr/bin/env python3
"""Generate Feedback-Mobile-Status-TD-Outline.docx — TD slide outline + Phase 4/5 checklists."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT = "docs/Feedback-Mobile-Status-TD-Outline.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_checks(doc, items):
    for item in items:
        doc.add_paragraph(f"☐  {item}", style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_slide(doc, number, title, body_fn):
    add_heading(doc, f"Slide {number} — {title}", 2)
    body_fn(doc)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Feedback Mobile Status — Technical Director Slide Outline", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Audience: ").bold = True
    p.add_run("Technical Director")
    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run("July 13, 2026")
    p = doc.add_paragraph()
    p.add_run("Scope: ").bold = True
    p.add_run(
        "Capacitor production app vs React Native (Expo) migration app; "
        "Phase 4/5 acceptance checklists included."
    )
    p = doc.add_paragraph()
    p.add_run("Product brand: ").bold = True
    p.add_run(
        "Feedback (com.feedbacklive.app). “Momentum” is the internal repo / DB codename."
    )
    p = doc.add_paragraph()
    p.add_run("Leave-behind tip: ").bold = True
    p.add_run(
        "Print Slide 3 + Phase 4/5 checklists; keep slides 5–9 as talking track."
    )

    # -------------------------------------------------------------------------
    add_heading(doc, "Part A — Slide Outline", 1)

    add_slide(
        doc,
        1,
        "Title",
        lambda d: add_bullets(
            d,
            [
                "Feedback Mobile: Capacitor vs RN Migration",
                "One line: Capacitor is production; RN is mid-migration "
                "(Phase 4 QA → Phase 5 harden → Phase 6 cutover).",
            ],
        ),
    )

    add_slide(
        doc,
        2,
        "Agenda",
        lambda d: add_bullets(
            d,
            [
                "Recommendation",
                "Shared backend",
                "Capacitor (prod) status",
                "Feedback RN status & phases",
                "Side-by-side",
                "Risks & decisions",
                "Phase 4 / 5 acceptance checklists",
            ],
        ),
    )

    def slide3(d):
        add_para(d, "Decision slide", bold=True)
        add_table(
            d,
            ["Keep", "Do next", "Don’t"],
            [
                (
                    "Ship/maintain Capacitor (com.feedbacklive.app)",
                    "Finish Phase 4 device QA on physical iPhone",
                    "Cut over without Phase 6 approval",
                ),
                (
                    "RN as isolated migration (…app.rn)",
                    "Define Phase 5 MVP gates",
                    "Demo RN as “App Store ready”",
                ),
            ],
        )
        add_para(d, "Ask of TD: ", bold=True)
        add_para(
            d,
            "Confirm Capacitor = prod; approve Phase 4/5 checklist as cutover gates.",
        )

    add_slide(doc, 3, "Recommendation", slide3)

    add_slide(
        doc,
        4,
        "Shared platform",
        lambda d: add_bullets(
            d,
            [
                "Cloudflare Worker + D1 + R2 + Durable Objects",
                "Hybrid auth (Mocha + Google/Apple/email)",
                "Same concert loop: feed → capture → upload → publish",
                "Backend strong; risk is client-native "
                "(capture, OAuth, upload, Android)",
            ],
        ),
    )

    def slide5(d):
        add_bullets(
            d,
            [
                "React SPA in Capacitor 7 WebView",
                "Live Workers URL → JS updates without new TestFlight (optional)",
                "iOS-first: Capgo camera, native audio, Apple/Google Sign-In, "
                "Photos, geo, push",
                "Full product surface (~37 routes): feed, discover, capture, "
                "queue, profile, admin, partners…",
                "Gaps: Android thinner; background upload stub; dual OAuth with RN",
            ],
        )
        add_para(d, "Status line: Production-capable web + iOS.", bold=True)

    add_slide(doc, 5, "Capacitor (production)", slide5)

    def slide6(d):
        add_bullets(
            d,
            [
                "Expo 57 + Expo Router · com.feedbacklive.app.rn · “Feedback RN”",
                "Isolated app (does not touch Capacitor ios/ / android/)",
                "Phases 1–4 coded · Phase 4 device QA pending · 5–6 open",
                "Core loop in code: auth → feeds → Vision Camera → outbox → "
                "Worker multipart",
            ],
        )
        add_para(
            d,
            "Status line: Credible parallel track; not cutover-ready.",
            bold=True,
        )

    add_slide(doc, 6, "Feedback RN (migration)", slide6)

    def slide7(d):
        add_bullets(
            d,
            [
                "[x] Phase 1 — Expo scaffold",
                "[x] Phase 2 — Auth / API / push / location",
                "[x] Phase 3 — Nav + read feeds",
                "[x] Phase 4 — Capture + upload ← device QA gate",
                "[ ] Phase 5 — Hardening",
                "[ ] Phase 6 — Cutover (explicit approval)",
            ],
        )

    add_slide(doc, 7, "Phase roadmap", slide7)

    def slide8(d):
        add_table(
            d,
            ["", "Capacitor", "RN"],
            [
                ("Production", "Yes", "No"),
                ("Breadth", "Full product", "Core loop + subset"),
                (
                    "Capture",
                    "Capgo + audio plugin",
                    "Vision Camera + audio-session",
                ),
                (
                    "Android",
                    "Thin WebView",
                    "Early (Google/audio iOS-only)",
                ),
                (
                    "Updates",
                    "Workers URL / bundled dist",
                    "Native rebuild / future EAS",
                ),
            ],
        )

    add_slide(doc, 8, "Side-by-side", slide8)

    add_slide(
        doc,
        9,
        "Risks",
        lambda d: add_bullets(
            d,
            [
                "Two apps, one brand (demo confusion)",
                "Dual OAuth audiences misconfigured",
                "RN capture unproven on hardware",
                "Android underinvested on both tracks",
                "Scope creep before cutover (don’t require full web parity)",
                "Media format quirks (e.g. .webm on AVFoundation)",
            ],
        ),
    )

    add_slide(
        doc,
        10,
        "Decisions needed",
        lambda d: add_bullets(
            d,
            [
                "Android: invest on Capacitor, hold, or wait for RN?",
                "RN MVP cutover surface (see checklist) vs “full parity later”",
                "Who signs Phase 4 / 5 / 6?",
                "Store strategy at Phase 6: replace listing vs new app?",
            ],
        ),
    )

    add_slide(
        doc,
        11,
        "Phase 4 acceptance (overview)",
        lambda d: add_bullets(
            d,
            [
                "Gate: Physical iPhone only · rebuild after native changes · "
                "pass = enter Phase 5",
                "Detail in Part B checklist",
            ],
        ),
    )

    add_slide(
        doc,
        12,
        "Phase 5 acceptance (overview)",
        lambda d: add_bullets(
            d,
            [
                "Gate: MVP hardening vs Capacitor · pass = eligible for "
                "Phase 6 proposal",
                "Detail in Part B checklist",
            ],
        ),
    )

    def slide13(d):
        add_para(d, "Fill in with owners before the meeting.", bold=True)
        add_table(
            d,
            ["Week", "Capacitor", "RN"],
            [
                ("W1", "OAuth / capture / upload polish", "Phase 4 device QA"),
                ("W2", "Android decision", "Start Phase 5 P0 items"),
            ],
        )

    add_slide(doc, 13, "Timeline / next 2 weeks", slide13)

    add_slide(
        doc,
        14,
        "Ask / close",
        lambda d: add_bullets(
            d,
            [
                "Approve checklists",
                "Confirm prod = Capacitor",
                "Schedule Phase 4 QA review date",
            ],
        ),
    )

    # -------------------------------------------------------------------------
    add_heading(doc, "Part B — Phase 4 / 5 Acceptance Checklists", 1)
    add_para(
        doc,
        "Use as a signed gate doc. Owner initials each row; TD signs the phase.",
    )

    add_heading(doc, "Phase 4 — Capture + upload (device QA)", 2)
    add_para(
        doc,
        "Gate: Physical iPhone only. Rebuild after native changes. "
        "Pass = enter Phase 5.",
        bold=True,
    )

    add_heading(doc, "Environment", 3)
    add_checks(
        doc,
        [
            "Physical iPhone (not simulator)",
            "Dev client rebuilt after Phase 4 natives (mobile:ios / expo run:ios)",
            "Points at intended Worker (EXPO_PUBLIC_API_BASE_URL)",
            "Signed in with Google and Apple (at least one full session each)",
        ],
    )

    add_heading(doc, "Auth / session", 3)
    add_checks(
        doc,
        [
            "Cold start restores session (bearer / SecureStore)",
            "/api/users/me succeeds after relaunch",
            "Sign out clears session; protected APIs fail until re-auth",
            "OAuth deep-link fallback path does not brick the app if native "
            "SDK path fails",
        ],
    )

    add_heading(doc, "Capture", 3)
    add_checks(
        doc,
        [
            "Camera + mic permission prompts show correct Feedback copy",
            "Rear camera preview starts reliably",
            "Record ≤60s produces MP4 under Documents/captures",
            "Audio present on recording (not silent / broken mux)",
            "Audio-session prepare/restore does not leave other app audio "
            "broken after exit",
            "Soft-fail if audio module missing is acceptable in that build "
            "(or module present and working — document which)",
            "Show-match HUD behaves reasonably with location granted "
            "(and degrades cleanly if denied)",
            "Abort / cancel mid-record does not leave a corrupt pending handoff",
        ],
    )

    add_heading(doc, "Handoff → upload", 3)
    add_checks(
        doc,
        [
            "Post-record navigates to review (/upload) with clip attached",
            "Caption / artist / venue / song fields save into outbox meta",
            "Share enqueues outbox item",
            "Queue tab lists item with clear status "
            "(pending / uploading / failed / done)",
            "Multipart: init → PUT parts → complete → poll until published",
            "Keep-awake holds during upload; screen can idle without killing "
            "the job (spot-check)",
            "Clip appears in feed / profile after publish",
            "Optional: saved to Photos via media library",
            "Kill app mid-upload → reopen → pending recovery works "
            "(boot and/or queue tab)",
            "Retry works on a forced failure (airplane mode then retry)",
        ],
    )

    add_heading(doc, "Smoke (non-blocking for Phase 4 — note failures)", 3)
    add_checks(
        doc,
        [
            "Home feed loads",
            "Clip player opens at least one known-good format",
            "Alerts list loads",
            "Discover / artist / venue open without crash",
        ],
    )

    add_heading(doc, "Phase 4 exit criteria", 3)
    add_checks(
        doc,
        [
            "All Capture + Handoff/upload checks pass on one primary iPhone "
            "OS version",
            "Known issues logged (format: ID, severity, workaround) — "
            "no Sev-1 blockers",
            "TD / eng lead signs Phase 4 → Phase 5 may start",
        ],
    )
    add_para(doc, "Sev-1 (blocks Phase 4):", bold=True)
    add_para(
        doc,
        "Cannot record; silent/unplayable upload; cannot authenticate; "
        "upload never publishes; crash loop on capture/queue.",
    )

    # Phase 5
    add_heading(doc, "Phase 5 — Hardening (pre-cutover)", 2)
    add_para(
        doc,
        "Split into P0 (required for Phase 6 proposal) and P1 "
        "(can defer post-cutover if TD agrees).",
    )

    add_heading(doc, "P0 — Must pass before proposing cutover", 3)

    add_para(doc, "Capture parity (MVP)", bold=True)
    add_checks(
        doc,
        [
            "Documented MVP vs Capacitor: what is in / out "
            "(AudD live ID, zoom/flip, sticky show session)",
            "Agreed MVP items implemented or explicitly waived in writing",
            "No Sev-1 capture/upload regressions vs Phase 4 baseline",
            "Upload metadata path acceptable for launch "
            "(manual fields OK if AudD deferred — must be explicit)",
        ],
    )

    add_para(doc, "Social / engagement", bold=True)
    add_checks(
        doc,
        [
            "Like + save wired in player / tiles (context mounted and used)",
            "State survives navigation; errors surfaced to user",
        ],
    )

    add_para(doc, "Push", bold=True)
    add_checks(
        doc,
        [
            "Permission → device token obtained",
            "Token POSTed to Worker and stored",
            "Test notification received on device "
            "(or documented staging substitute)",
        ],
    )

    add_para(doc, "Alerts / navigation", bold=True)
    add_checks(
        doc,
        [
            "Tap notification / alert opens the correct clip or entity "
            "(or agreed subset)",
            "Discover respects search query from home "
            "(or home no longer promises it)",
        ],
    )

    add_para(doc, "Auth / dual-client", bold=True)
    add_checks(
        doc,
        [
            "Worker secrets verified: Capacitor + RN Google/Apple audiences",
            "Capacitor prod auth still green after RN secret changes "
            "(regression)",
        ],
    )

    add_para(doc, "Stability / media", bold=True)
    add_checks(
        doc,
        [
            "Playback policy for unsupported formats (e.g. .webm): "
            "skip, transcode, or hide — implemented and tested",
            "Crash-free soak: 20+ capture→upload cycles, 0 Sev-1",
        ],
    )

    add_para(doc, "Android (decide one)", bold=True)
    add_checks(
        doc,
        [
            "Option A: Android out of Phase 6 scope (iOS-only cutover) — "
            "written waiver",
            "Option B: Android Google Sign-In + capture/upload smoke pass",
        ],
    )

    add_para(doc, "Release plumbing", bold=True)
    add_checks(
        doc,
        [
            "EAS (or equivalent) project wired if OTA/push/CI required "
            "for launch",
            "Bundle ID / display name plan for Phase 6 documented "
            "(keep .rn vs switch to prod ID)",
        ],
    )

    add_heading(doc, "P1 — Strongly preferred; defer only with TD waiver", 3)
    add_checks(
        doc,
        [
            "Live AudD while recording",
            "Zoom-while-record / camera flip",
            "Capgo-equivalent settle timings / multi-clip sticky show session",
            "Content-feed classify + ACR at upload",
            "Comments",
            "Other-user profiles / onboarding / email auth",
            "Admin / premium / partner surfaces (usually stay Capacitor/web)",
        ],
    )

    add_heading(doc, "Phase 5 exit criteria", 3)
    add_checks(
        doc,
        [
            "All P0 checked or waived with named owner + date",
            "P1 list reviewed; deferred items ticketed",
            "Capacitor remains healthy in parallel "
            "(no collision on native projects)",
            "Written Phase 6 cutover proposal attached "
            "(store strategy, rollback, feature freeze)",
            "TD signs Phase 5 → Phase 6 scheduling allowed",
        ],
    )

    add_heading(doc, "Phase 6 reminder (not a full checklist)", 2)
    add_para(
        doc,
        "Cutover only after Phase 5 sign-off: store listing plan, feature freeze, "
        "rollback (keep Capacitor live), dual-OAuth cleanup, explicit "
        "TD/product approval.",
    )

    add_heading(doc, "Sign-off", 2)
    add_table(
        doc,
        ["Phase", "Owner", "Date", "TD signature"],
        [
            ("Phase 4", "", "", ""),
            ("Phase 5", "", "", ""),
            ("Phase 6 (cutover approval)", "", "", ""),
        ],
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
