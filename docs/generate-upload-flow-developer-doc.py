#!/usr/bin/env python3
"""Generate upload-flow-developer-summary.docx — technical upload architecture reference."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "docs/upload-flow-developer-summary.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    return p


def main():
    doc = Document()

    title = doc.add_heading("Clip Upload Flow — Developer Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Document purpose: ").bold = True
    p.add_run("Technical reference for the resilient clip upload / recovery pipeline.")
    p = doc.add_paragraph()
    p.add_run("Last updated: ").bold = True
    p.add_run("June 17, 2026")

    doc.add_paragraph()

    # Overview
    add_heading(doc, "Overview", 1)
    add_para(
        doc,
        "Clip uploads use a client-side outbox + server-side multipart pipeline. "
        "The user taps Share and can leave immediately; upload continues in the background, "
        "survives refresh/offline, and retries automatically.",
    )
    add_code_block(
        doc,
        "Share → Outbox Queue → Multipart R2 Upload → Publish (R2 playback) → Stream Ingest → Feed\n"
        "         ↑ IDB + memory pin",
    )

    # Client Architecture
    add_heading(doc, "Client Architecture", 1)

    add_heading(doc, "Entry point", 2)
    add_para(doc, "All Share paths go through ClipUploadQueueContext via UploadClip.tsx → enqueue().")

    add_heading(doc, "Outbox queue (ClipUploadQueueContext.tsx)", 2)
    add_bullets(
        doc,
        [
            "FIFO: one clip uploads at a time",
            "States: queued → classifying → uploading → completing → processing → published (or paused / failed)",
            "Auto-retry with exponential backoff (5s → 120s max) for transient errors",
            "Wake lock while uploading (best-effort on mobile browsers)",
            "Global UI: ClipUploadStatusBanner.tsx",
        ],
    )

    add_heading(doc, "Local persistence (3 layers)", 2)
    add_table(
        doc,
        ["Layer", "File", "Purpose"],
        [
            ("Memory pin", "clip-blob-registry.ts", "In-tab blob survives IDB failures"),
            ("Memory cache", "blob-store.ts", "Fast lookup while IDB write is in flight"),
            ("IndexedDB", "idb.ts", "Survives refresh — meta in meta store, blobs in blobs store"),
        ],
    )
    add_para(
        doc,
        "On enqueue, video is written to IndexedDB before upload starts. "
        "Gallery save and thumbnail generation run in the background (background-persist.ts).",
    )

    add_heading(doc, "Job runner (runner.ts)", 2)
    add_para(doc, "Orchestrates one outbox job:")
    add_bullets(
        doc,
        [
            "Optional content classification (skipped when BYPASS_CONTENT_FEED_BIFURCATION = true)",
            "POST /api/uploads/init — creates draft clip + upload session",
            "Multipart upload of video chunks",
            "Thumbnail attach + POST /api/uploads/:id/complete",
            "Poll until server reports published",
            "Clear local outbox on success",
        ],
    )

    add_heading(doc, "Multipart upload (multipart-upload.ts)", 2)
    add_bullets(
        doc,
        [
            "Chunks: 5 MB parts (UPLOAD_PART_SIZE_BYTES)",
            "Worker mode only — parts go PUT /api/uploads/:sessionId/parts/:n → Worker → R2 "
            "(direct R2 presign disabled due to CORS)",
            "Resume: fetches session status, skips already-uploaded parts",
            "Session invalid (404/409): clears session, re-inits, retries once",
            "Refresh recovery: if blob is missing but all parts are on the server, completes via session without local video",
        ],
    )

    add_heading(doc, "Refresh / offline recovery", 2)
    add_para(doc, "On app load, queue hydrates from IndexedDB meta:")
    add_bullets(
        doc,
        [
            "In-flight states reset to queued",
            "Blobs loaded from IDB → memory pin",
            "If blob missing but sessionId exists → check server status; finish or resume",
            "If blob missing and session incomplete → wait up to 45s for IDB, then fail clearly",
            "Paused jobs retry on timer; poll loop does not bypass blob-wait backoff",
        ],
    )

    # Server Architecture
    add_heading(doc, "Server Architecture", 1)

    add_heading(doc, "Database (migrations/60.sql)", 2)
    add_bullets(
        doc,
        [
            "upload_sessions — multipart state, completed parts JSON, expiry",
            "clips.upload_status — uploading | uploaded | processing | ready | failed",
            "clips.r2_raw_key — R2 object key for raw video",
        ],
    )

    add_heading(doc, "API endpoints (upload-endpoints.ts)", 2)
    add_table(
        doc,
        ["Endpoint", "Role"],
        [
            ("POST /api/uploads/init", "Create draft clip (is_draft=1, video_url='pending:upload') + R2 multipart session"),
            ("PUT .../parts/:n", "Upload one chunk to R2, track etag in D1"),
            ("POST .../complete", "Finalize multipart, publish clip with R2 playback URL, trigger Stream ingest"),
            ("GET .../status", "Progress, clipPublished, completed part numbers"),
            ("POST .../thumbnail", "Attach client JPEG early for grid display"),
        ],
    )

    add_heading(doc, "Publish pipeline (upload-processor.ts)", 2)
    add_para(doc, "On complete:", bold=True)
    add_bullets(
        doc,
        [
            "Set video_url = /api/files/{r2Key}, is_draft=0, status='published'",
            "waitUntil(processClipStreamIngestById) — immediate Stream ingest (not cron-only)",
            "Award points + notify user + broadcast feed update",
        ],
    )
    add_para(doc, "Stream ingest (also runs on cron):", bold=True)
    add_bullets(
        doc,
        [
            "Cloudflare Stream /copy from {PUBLIC_APP_URL}/api/files/{key}",
            "On success: set Stream URLs, upload_status='ready'",
            "On failure: keep R2 playback live, retry on next cron",
        ],
    )

    add_heading(doc, "Playback resolution (clip-playback.ts)", 2)
    add_bullets(
        doc,
        [
            "Feed tiles: Stream MP4 when available, else R2 /api/files/...",
            "Modal: Stream HLS when available, else R2 direct",
            "Placeholder (pending:upload): falls back to r2_raw_key if present",
        ],
    )

    add_heading(doc, "Feed visibility", 2)
    add_bullets(
        doc,
        [
            "My clips: includes drafts and in-progress uploads",
            "Latest clips: requires is_draft=0 (published after R2 complete)",
        ],
    )

    # End-to-End Sequence
    add_heading(doc, "End-to-End Sequence", 1)
    add_code_block(
        doc,
        "User taps Share\n"
        "  └─ enqueue(payload)\n"
        "       ├─ registerClipBlob + cacheOutboxBlobs\n"
        "       ├─ await persistOutboxVideo (IndexedDB)\n"
        "       └─ processNext()\n"
        "\n"
        "processNext → runOutboxJob\n"
        "  ├─ POST /api/uploads/init\n"
        "  │    └─ INSERT draft clip + upload_sessions row + R2 CreateMultipartUpload\n"
        "  ├─ PUT /api/uploads/:id/parts/1..N  (resume-aware)\n"
        "  ├─ POST /api/uploads/:id/thumbnail  (optional, early)\n"
        "  ├─ POST /api/uploads/:id/complete\n"
        "  │    └─ R2 CompleteMultipartUpload\n"
        "  │    └─ publishClipWithR2Playback (is_draft=0, video_url set)\n"
        "  │    └─ waitUntil → Stream ingest\n"
        "  └─ GET /api/uploads/:id/status (poll until clipPublished)\n"
        "\n"
        "Client: status → published, clear outbox, dismiss banner",
    )

    # Key Files
    add_heading(doc, "Key Files", 1)
    add_table(
        doc,
        ["Area", "Files"],
        [
            ("Queue / UI", "ClipUploadQueueContext.tsx, ClipUploadStatusBanner.tsx"),
            ("Outbox", "runner.ts, multipart-upload.ts, blob-store.ts, idb.ts, clip-blob-registry.ts"),
            ("Server upload", "upload-endpoints.ts, upload-clip-create.ts"),
            ("Processing", "upload-processor.ts, scheduled.ts"),
            ("Playback", "clip-playback.ts, clip-poster-url.ts, r2-serve.ts"),
            ("Shared", "shared/upload.ts (part size, types)"),
        ],
    )

    # Operational Notes
    add_heading(doc, "Operational Notes", 1)
    add_bullets(
        doc,
        [
            "Deploy: npm run build && wrangler deploy (Worker + frontend)",
            "Env: PUBLIC_APP_URL must be set for Stream to fetch R2 files",
            "Rate limit: 400 requests/hour shared bucket for all multipart endpoints per user",
            "Capacitor / native background upload: scaffolded (native-bridge.ts, gallery-save.ts) "
            "but not fully wired — browser relies on in-tab queue + IDB recovery",
        ],
    )

    # Design Principles
    add_heading(doc, "Design Principles", 1)
    add_bullets(
        doc,
        [
            "Never lose the clip locally — memory pin + IDB before network",
            "Resume, don't restart — part-level resume + session recovery after refresh",
            "Publish fast, enhance later — R2 playback immediately; Stream is an upgrade",
            "Fail soft, retry hard — transient errors → paused + backoff; permanent errors → failed with dismiss",
            "One clip at a time — avoids bandwidth contention on venue Wi-Fi",
        ],
    )

    add_para(
        doc,
        "See also: docs/upload-recovery-outline.md (plain-language plan), "
        "docs/upload-outbox-build.md (original build sketch).",
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("Momentum — Clip Upload Flow (Developer Summary)")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
