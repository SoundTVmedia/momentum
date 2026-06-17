#!/usr/bin/env python3
"""Generate upload-recovery-outline.docx from structured content."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = "docs/upload-recovery-outline.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Momentum — Upload & Recovery Plan", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Document purpose: ").bold = True
    p.add_run(
        "Plain-language outline of how we're making clip uploads reliable at live shows."
    )
    p = doc.add_paragraph()
    p.add_run("Last updated: ").bold = True
    p.add_run("June 17, 2026")

    doc.add_paragraph()

    # Problem
    add_heading(doc, "The Problem Today", 1)
    add_para(
        doc,
        'When someone records a clip at a show and taps Share:',
    )
    add_bullets(
        doc,
        [
            'The video has to finish uploading before anything is truly "saved."',
            "If they close the app, lose signal, or switch screens, the upload can fail and the clip may be lost.",
            "On phones, uploads only work reliably while the app stays open in the foreground.",
            'There\'s no "your clip is live" notification when processing finishes.',
            'The server doesn\'t really know a clip is "on the way" until the entire upload completes.',
        ],
    )
    add_para(
        doc,
        "In short: Recording works, but getting the clip safely to the feed in bad venue Wi-Fi is fragile.",
        bold=True,
    )

    # Goal
    add_heading(doc, "The Goal", 1)
    add_para(doc, "Make posting feel like Instagram or TikTok at a concert:")
    add_bullets(
        doc,
        [
            "Tap Share → leave immediately (back to feed, record another clip).",
            'See "Uploading…" on your clips even if you close and reopen the app.',
            "Pick up where you left off if the connection drops.",
            "Get a notification when the clip is ready.",
            "Optionally save a copy to your camera roll.",
        ],
    )

    # Big pieces
    add_heading(doc, "The Big Pieces (What We're Adding)", 1)

    add_heading(doc, "1. A Real Phone App Wrapper (Capacitor)", 2)
    add_para(
        doc,
        "We wrap the existing web app in Capacitor so it behaves more like a native phone app.",
    )
    add_table(
        doc,
        ["Feature", "What it means for users"],
        [
            ("Native video recording", "More reliable camera; video saved to the phone's storage right away"),
            ("Background upload", "Upload keeps going (or resumes) when you switch apps or lock the screen"),
            ("Save to gallery", '"Also save to Photos" after recording'),
            ("Push notifications", '"Your clip from [Artist] at [Venue] is live!"'),
        ],
    )
    add_para(doc, "The website still works in the browser; the app gets the extra reliability.")

    add_heading(doc, "2. Split the Video Into Pieces Before Sending", 2)
    add_para(
        doc,
        "A 60-second clip is still a big file on spotty Wi-Fi. Instead of sending it in one shot:",
    )
    add_bullets(
        doc,
        [
            "The app cuts it into small chunks (a few megabytes each).",
            "Each chunk uploads on its own.",
            "If chunk 3 of 6 fails, we only retry chunk 3 — not the whole video.",
        ],
    )
    add_para(
        doc,
        "Analogy: Mailing a thick book by sending a few pages at a time. If one envelope gets lost, you resend only that envelope.",
        bold=True,
    )

    add_heading(doc, "3. Upload Straight to Cloud Storage (Not Through Our Server)", 2)
    add_para(
        doc,
        "Chunks go directly to Cloudflare R2 (our file storage), using temporary secure links.",
    )
    add_para(doc, "Why it matters:", bold=True)
    add_bullets(
        doc,
        [
            "Faster uploads (no middleman).",
            "Our server doesn't get overloaded holding big video files.",
            "Easier to resume — we know which chunks already landed.",
        ],
    )

    add_heading(doc, "4. Save Progress in Two Places", 2)
    add_para(doc, "On the phone (local queue):", bold=True)
    add_bullets(
        doc,
        [
            "Which clips are waiting to upload.",
            "Which chunks finished.",
            "Caption, artist, venue, etc.",
        ],
    )
    add_para(
        doc,
        "So if the app crashes or you close it, reopening picks up where you left off.",
    )
    add_para(doc, "In the database (server):", bold=True)
    add_bullets(
        doc,
        [
            'A clip row is created as soon as you tap Share — status: "uploading."',
            'When all chunks arrive: status → "uploaded."',
            'When processing finishes: status → "live."',
        ],
    )
    add_para(
        doc,
        "Analogy: You get a receipt the moment you order (Share), not only when the food arrives (upload done).",
        bold=True,
    )

    add_heading(doc, "5. A Simple Background Worker (No Fancy Queue System)", 2)
    add_para(doc, 'A small scheduled job runs about every minute and asks:')
    q = doc.add_paragraph()
    q.add_run('"Are there any clips that finished uploading but aren\'t processed yet?"').italic = True
    add_para(doc, "For each one it:")
    add_bullets(
        doc,
        [
            "Prepares the video for playback (e.g. send to Cloudflare Stream).",
            "Makes a thumbnail.",
            "Runs any other steps (music ID, content checks, etc.).",
            "Marks the clip published and sends a push notification.",
        ],
    )
    add_para(
        doc,
        "We're not adding Redis or a heavy job system for MVP — a database table plus a cron job is enough to start.",
    )

    # User experience
    add_heading(doc, "Step-by-Step: What the User Experiences", 1)

    add_heading(doc, "Recording", 2)
    add_bullets(
        doc,
        [
            "User opens the app and records up to 60 seconds.",
            "Video is saved on the device immediately.",
            "Optional: copy saved to camera roll.",
        ],
    )

    add_heading(doc, "Caption Screen", 2)
    add_bullets(
        doc,
        [
            "User adds caption, artist, venue (often auto-filled from location).",
            "Taps Share.",
        ],
    )

    add_heading(doc, "Right After Share (within a second)", 2)
    add_bullets(
        doc,
        [
            'App creates an "uploading" clip on the server.',
            "Video is queued locally on the phone.",
            "User can go home, record another clip, or close the app.",
            '"Uploading…" tile shows on My Clips with progress %.',
        ],
    )

    add_heading(doc, "Upload (in background)", 2)
    add_bullets(
        doc,
        [
            "App uploads chunks one by one to cloud storage.",
            "Progress updates (e.g. 40% → 80%).",
            "If Wi-Fi drops: pause, retry when back online from the last good chunk.",
        ],
    )

    add_heading(doc, "Processing (on the server)", 2)
    add_bullets(
        doc,
        [
            "When all chunks are in, server assembles the full video.",
            "Background worker processes it (playback, thumbnail, etc.).",
            "Clip goes live in the feed.",
        ],
    )

    add_heading(doc, "Done", 2)
    add_bullets(
        doc,
        [
            "User gets a push notification.",
            '"Uploading…" tile becomes a normal clip in My Clips and the feed.',
        ],
    )

    # Error handling
    add_heading(doc, "What Happens When Things Go Wrong", 1)
    add_table(
        doc,
        ["Situation", "What happens"],
        [
            ("Bad Wi-Fi mid-upload", "Pause, resume from last successful chunk"),
            ("App closed", "Local queue + server both remember the clip; resume on reopen"),
            ("Upload fails completely", '"Failed — tap to retry" on My Clips'),
            ("User cancels", "Remove from queue; delete unfinished server record"),
            ("Processing fails on server", 'Worker retries; user may see "processing" longer or get a retry'),
        ],
    )

    # Summary
    add_heading(doc, "What We're Changing (Summary)", 1)

    add_heading(doc, "For Users", 2)
    add_bullets(
        doc,
        [
            "Share and leave — no waiting on a spinner.",
            "See uploading clips in My Clips.",
            "Uploads survive bad signal and app restarts (especially in the native app).",
            "Notifications when clips go live.",
            "Optional save to Photos.",
        ],
    )

    add_heading(doc, "Behind the Scenes", 2)
    add_table(
        doc,
        ["Area", "Change"],
        [
            ("Phone app", "Capacitor wrapper + native recording + background upload"),
            ("Upload method", "Chunked, resumable, direct to cloud storage"),
            ("Local storage", "Persistent queue on device (survives refresh/restart)"),
            ("Database", "Clip + upload status tracked from Share, not only after upload"),
            ("Background job", "Minute-by-minute worker processes finished uploads"),
            ("Notifications", "Push when clip is ready"),
        ],
    )

    add_heading(doc, "What We're NOT Doing (For Now)", 2)
    add_bullets(
        doc,
        [
            "Redis or a complex job queue.",
            "Uploading after the browser tab is fully killed on web only (native app handles that).",
            "Showing half-uploaded clips in the public feed (only on My Clips until live).",
        ],
    )

    # Build order
    add_heading(doc, "Build Order (How We'll Roll It Out)", 1)

    phases = [
        (
            "Phase 1 — Server foundation",
            "Database tables, secure upload links, start/finish upload APIs, background processor.",
        ),
        (
            "Phase 2 — Web app queue",
            "Tap Share → queue locally + create server record; show progress on My Clips; retry on failure.",
        ),
        (
            "Phase 3 — Native app (Capacitor)",
            "App store build, native camera, background upload, push notifications, save to gallery.",
        ),
        (
            "Phase 4 — Polish",
            "Better error messages, cleanup of abandoned uploads, analytics, tuning for slow networks.",
        ),
    ]
    for title, desc in phases:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(f"\n{desc}")

    doc.add_paragraph()

    # One sentence
    add_heading(doc, "One-Sentence Summary", 1)
    summary = doc.add_paragraph()
    s = summary.add_run(
        'We\'re turning "wait on a loading screen until the whole video uploads" into '
        '"tap Share, leave immediately, and we\'ll reliably get your clip live in the background '
        '— even at a show with bad Wi-Fi."'
    )
    s.bold = True
    s.italic = True

    doc.add_paragraph()
    footer = doc.add_paragraph("Momentum — Upload & Recovery Plan")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
